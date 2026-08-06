"""发行体清单读取（CSV / DOCX）。"""

from __future__ import annotations

import csv
import logging
import re
from pathlib import Path

from docx import Document

logger = logging.getLogger(__name__)

_NAME_HEADERS = {
    "发行体",
    "issuer",
    "issuer_name",
    "name",
    "entity",
    "主体",
    "发行人",
}


def _norm_header(h: str) -> str:
    return re.sub(r"\s+", "", (h or "").strip().lower())


def _pick_name_column(headers: list[str]) -> int:
    normalized = [_norm_header(h) for h in headers]
    targets = {_norm_header(x) for x in _NAME_HEADERS}
    for i, h in enumerate(normalized):
        if h in targets or "issuer" in h or "发行" in h:
            return i
    return 0


def load_issuers_from_csv(path: Path) -> list[str]:
    names: list[str] = []
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        sample = f.read(4096)
        f.seek(0)
        try:
            dialect = csv.Sniffer().sniff(sample, delimiters=",\t;")
        except csv.Error:
            dialect = csv.excel
        reader = csv.reader(f, dialect)
        rows = list(reader)
    if not rows:
        return []
    first = rows[0]
    targets = {_norm_header(x) for x in _NAME_HEADERS}
    has_header = any(_norm_header(c) in targets for c in first)
    start = 1 if has_header else 0
    col = _pick_name_column(first) if has_header else 0
    for row in rows[start:]:
        if not row or col >= len(row):
            continue
        name = (row[col] or "").strip()
        if name and name not in names:
            names.append(name)
    return names


def load_issuers_from_docx(path: Path) -> list[str]:
    doc = Document(str(path))
    names: list[str] = []
    targets = {_norm_header(x) for x in _NAME_HEADERS}

    def _add(text: str) -> None:
        t = (text or "").strip()
        if not t or t in names:
            return
        if _norm_header(t) in targets:
            return
        names.append(t)

    for table in doc.tables:
        for i, row in enumerate(table.rows):
            cells = [(c.text or "").strip() for c in row.cells]
            if not any(cells):
                continue
            if i == 0 and any(_norm_header(c) in targets for c in cells):
                continue
            for c in cells:
                if c:
                    _add(c)
                    break

    if not names:
        for p in doc.paragraphs:
            line = (p.text or "").strip()
            if line:
                _add(line)
    return names


def resolve_input_file(input_dir: Path, candidates: list[str]) -> Path:
    for name in candidates:
        p = input_dir / name
        if p.is_file():
            return p
    for pattern in ("*.csv", "*.docx"):
        found = sorted(input_dir.glob(pattern))
        if found:
            return found[0]
    raise FileNotFoundError(
        f"未在 {input_dir} 找到发行体清单（期望: {', '.join(candidates)}）"
    )


def load_issuer_records(
    input_dir: Path, candidates: list[str] | None = None
) -> tuple[Path, list[dict[str, str]]]:
    """读取发行体清单，返回 (路径, [{name, category}, ...])。"""
    candidates = candidates or ["issuer_list.csv", "issuer_list.docx"]
    path = resolve_input_file(input_dir, candidates)
    suffix = path.suffix.lower()
    if suffix == ".csv":
        records = _load_records_from_csv(path)
    elif suffix in {".docx", ".doc"}:
        names = load_issuers_from_docx(path)
        records = [{"name": n, "category": ""} for n in names]
    else:
        raise ValueError(f"不支持的清单格式: {path.name}")
    logger.info("已加载发行体记录 %d 条，来源: %s", len(records), path)
    return path, records


def load_issuer_list(
    input_dir: Path, candidates: list[str] | None = None
) -> tuple[Path, list[str]]:
    """读取发行体清单，返回 (文件路径, 去重后的名称列表)。"""
    path, records = load_issuer_records(input_dir, candidates)
    return path, [r["name"] for r in records]


def _load_records_from_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        sample = f.read(4096)
        f.seek(0)
        try:
            dialect = csv.Sniffer().sniff(sample, delimiters=",\t;")
        except csv.Error:
            dialect = csv.excel
        reader = csv.reader(f, dialect)
        rows = list(reader)
    if not rows:
        return []
    first = rows[0]
    targets = {_norm_header(x) for x in _NAME_HEADERS}
    has_header = any(_norm_header(c) in targets for c in first)
    name_col = _pick_name_column(first) if has_header else 0
    cat_col = -1
    if has_header:
        for i, h in enumerate(first):
            hn = _norm_header(h)
            if hn in {"分类", "category", "类别", "债券分类"} or "分类" in h:
                cat_col = i
                break
    start = 1 if has_header else 0
    out: list[dict[str, str]] = []
    seen: set[str] = set()
    for row in rows[start:]:
        if not row or name_col >= len(row):
            continue
        name = (row[name_col] or "").strip()
        if not name or name in seen:
            continue
        seen.add(name)
        cat = ""
        if cat_col >= 0 and cat_col < len(row):
            cat = (row[cat_col] or "").strip()
        out.append({"name": name, "category": cat})
    return out
