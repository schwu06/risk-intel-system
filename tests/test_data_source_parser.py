from __future__ import annotations

import tempfile
import unittest
from pathlib import Path
from unittest.mock import patch

from docx import Document
from openpyxl import Workbook

from app.services.data_source_parser import parse_uploaded_file
from app.services.industry_analysis import report_json_to_html


class _FakePdfPage:
    def __init__(self, text: str) -> None:
        self._text = text

    def extract_text(self) -> str:
        return self._text


class _FakePdfReader:
    def __init__(self, _path: str) -> None:
        self.pages = [
            _FakePdfPage("第一页包含足够长度的原生文本，用于验证页码标记会被保留，并避免进入OCR分支。"),
            _FakePdfPage("第二页也包含足够长度的原生文本，用于验证不同页面不会被混淆，并避免进入OCR分支。"),
        ]


class LegacyParserRegressionTests(unittest.TestCase):
    def test_txt_upload_and_parse_returns_string(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "sample.txt"
            path.write_text("第一段\n\n第二段 2025年 12.5%", encoding="utf-8")
            parsed = parse_uploaded_file(path)
        self.assertIsInstance(parsed, str)
        self.assertIn("第二段 2025年 12.5%", parsed)

    def test_pdf_page_markers_are_preserved(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "sample.pdf"
            path.write_bytes(b"%PDF-test")
            with patch("pypdf.PdfReader", _FakePdfReader):
                parsed = parse_uploaded_file(path)
        self.assertIn("[PDF 第 1 页]", parsed)
        self.assertIn("[PDF 第 2 页]", parsed)

    def test_docx_paragraphs_and_tables_are_parsed(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "sample.docx"
            doc = Document()
            doc.add_paragraph("段落内容")
            table = doc.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "指标"
            table.cell(0, 1).text = "120亿元"
            doc.save(path)
            parsed = parse_uploaded_file(path)
        self.assertIn("段落内容", parsed)
        self.assertIn("指标 | 120亿元", parsed)

    def test_xlsx_sheet_and_numeric_value_are_parsed(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "sample.xlsx"
            wb = Workbook()
            ws = wb.active
            ws.title = "财务数据"
            ws.append(["收入", 123456789.125])
            wb.save(path)
            wb.close()
            parsed = parse_uploaded_file(path)
        self.assertIn("[工作表: 财务数据]", parsed)
        self.assertIn("123456789.125", parsed)

    def test_report_json_to_html_is_unchanged(self) -> None:
        html = report_json_to_html(
            {
                "title": "测试报告",
                "summary": "摘要",
                "sections": [{"heading": "行业概况", "content": "正文"}],
                "risk_outlook": "风险展望正文",
            }
        )
        self.assertIn("<h1 class=\"report-title\">测试报告</h1>", html)
        self.assertIn("<h2>执行摘要</h2>", html)
        self.assertIn("<h2>行业概况</h2>", html)
        self.assertIn("<h2>风险展望</h2>", html)


if __name__ == "__main__":
    unittest.main()
