"""python-docx 中文日报导出。"""

from __future__ import annotations

from datetime import date, datetime
from io import BytesIO
from pathlib import Path
from typing import Any, Iterable, Optional
from urllib.parse import urlparse

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from app.config import MODULE_CODES
from app.database.models import CreditUpdate, DailyRiskEntry, EntityRisk, IndustryReport, TargetEntity
from app.services.chart_generator import extract_chart_specs, render_chart_png
from app.services.citation_validation import CITATION_RE

# 24小时核心新闻情报汇总 — 视觉规范色
COLOR_PRIMARY = RGBColor(0x1B, 0x36, 0x5D)  # 深蓝标题/分割线
COLOR_TITLE = RGBColor(0x2C, 0x3E, 0x50)  # 新闻标题
COLOR_META = RGBColor(0x7F, 0x8C, 0x8D)  # 元信息/副标题灰
COLOR_BODY = RGBColor(0x33, 0x33, 0x33)
COLOR_DIVIDER = "BDC3C7"


def _set_run_font(
    run,
    size_pt: float = 11,
    bold: bool = False,
    italic: bool = False,
    color: Optional[RGBColor] = None,
):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    sizes = {1: 18, 2: 14, 3: 12}
    _set_run_font(run, size_pt=sizes.get(level, 12), bold=True, color=COLOR_PRIMARY)


def _add_body(doc: Document, text: str, indent: bool = False) -> None:
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    _set_run_font(run, size_pt=11, color=COLOR_BODY)


def _add_meta_line(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}：")
    _set_run_font(r1, bold=True, color=RGBColor(0x44, 0x44, 0x44))
    r2 = p.add_run(value)
    _set_run_font(r2, color=COLOR_BODY)


def _add_light_divider(doc: Document) -> None:
    """浅灰色细分割线（段后间距约 6pt）。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), COLOR_DIVIDER)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _source_label(entry: DailyRiskEntry) -> str:
    if entry.source_title and entry.source_title.strip():
        return entry.source_title.strip()
    if entry.source_url:
        host = urlparse(entry.source_url).netloc
        if host:
            return host.removeprefix("www.")
    return "未知来源"


def _format_published_at(entry: DailyRiskEntry, fallback_date: date) -> str:
    if entry.published_at:
        return entry.published_at.strftime("%Y-%m-%d %H:%M")
    return fallback_date.isoformat()


def _category_tag(entry: DailyRiskEntry, module_name: str) -> str:
    parts: list[str] = []
    if entry.pillar_or_topic:
        parts.append(entry.pillar_or_topic)
    elif entry.risk_category:
        parts.append(entry.risk_category)
    if module_name and module_name not in parts:
        parts.append(module_name)
    return " / ".join(parts) if parts else "未分类"


def build_daily_report_docx(
    entries: Iterable[DailyRiskEntry],
    report_date: date,
    institution_name: str = "风险管理部",
    module_codes: Optional[list[str]] = None,
) -> Document:
    """生成《24小时核心新闻情报汇总》（仅标题/元数据/概要，不含风险研判与图表）。"""
    _ = institution_name  # 新版版式不再展示机构署名

    doc = Document()
    section = doc.sections[0]
    # 标准 Word 页边距：上下 2.54cm，左右 3.18cm
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    if module_codes:
        codes = [c.upper() for c in module_codes if c.upper() in MODULE_CODES]
        module_map = {c: MODULE_CODES[c] for c in codes}
    else:
        module_map = dict(MODULE_CODES)

    items = [e for e in entries if e.module_code in module_map]
    generated_at = datetime.now().strftime("%Y-%m-%d %H:%M")

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(4)
    tr = title_p.add_run("24小时核心新闻情报汇总")
    _set_run_font(tr, size_pt=22, bold=True, color=COLOR_PRIMARY)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(10)
    sr = sub.add_run(f"生成时间：{generated_at}  |  资讯总量：{len(items)} 条")
    _set_run_font(sr, size_pt=9, color=COLOR_META)

    # 标题下主色细分割线
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(0)
    rule.paragraph_format.space_after = Pt(12)
    rule_pr = rule._p.get_or_add_pPr()
    rule_bdr = OxmlElement("w:pBdr")
    rule_bottom = OxmlElement("w:bottom")
    rule_bottom.set(qn("w:val"), "single")
    rule_bottom.set(qn("w:sz"), "12")
    rule_bottom.set(qn("w:space"), "1")
    rule_bottom.set(qn("w:color"), "1B365D")
    rule_bdr.append(rule_bottom)
    rule_pr.append(rule_bdr)

    if not items:
        empty = doc.add_paragraph()
        empty.alignment = WD_ALIGN_PARAGRAPH.CENTER
        er = empty.add_run("本日暂无核心资讯条目。")
        _set_run_font(er, size_pt=10.5, color=COLOR_META)
        return doc

    for idx, e in enumerate(items, start=1):
        module_name = module_map.get(e.module_code, e.module_code)

        # 【序号与标题】小四 12pt 加粗 #2C3E50
        title_line = doc.add_paragraph()
        title_line.paragraph_format.space_before = Pt(4)
        title_line.paragraph_format.space_after = Pt(2)
        tr_item = title_line.add_run(f"{idx}. {e.title}")
        _set_run_font(tr_item, size_pt=12, bold=True, color=COLOR_TITLE)

        # 【元数据行】五号 10.5pt 斜体浅灰：来源 / 发布时间 / 分类标签
        meta = doc.add_paragraph()
        meta.paragraph_format.space_before = Pt(0)
        meta.paragraph_format.space_after = Pt(2)
        meta_text = (
            f"来源：{_source_label(e)}  |  "
            f"发布时间：{_format_published_at(e, report_date)}  |  "
            f"分类标签：{_category_tag(e, module_name)}"
        )
        mr = meta.add_run(meta_text)
        _set_run_font(mr, size_pt=10.5, italic=True, color=COLOR_META)

        # 数据源网络连接
        link_p = doc.add_paragraph()
        link_p.paragraph_format.space_before = Pt(0)
        link_p.paragraph_format.space_after = Pt(4)
        if e.source_url:
            lr = link_p.add_run(f"数据源网络连接：{e.source_url}")
        else:
            lr = link_p.add_run("数据源网络连接：暂无")
        _set_run_font(lr, size_pt=10.5, italic=True, color=COLOR_META)

        # 【核心概要】五号 10.5pt，1.25 倍行距，首行缩进 2 字符，两端对齐
        summary = doc.add_paragraph()
        summary.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        summary.paragraph_format.line_spacing = 1.25
        summary.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        summary.paragraph_format.first_line_indent = Pt(21)  # 10.5pt × 2 字符
        summary.paragraph_format.space_before = Pt(0)
        summary.paragraph_format.space_after = Pt(2)
        sr_body = summary.add_run(e.summary or "")
        _set_run_font(sr_body, size_pt=10.5, color=COLOR_BODY)

        _add_light_divider(doc)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(8)
    fr = footer.add_run("— 内部资料，未经许可不得对外传播 —")
    _set_run_font(fr, size_pt=9, color=COLOR_META)

    return doc

def _add_cited_body(doc: Document, text: str, citation_context: dict[str, Any]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    cursor = 0
    for match in CITATION_RE.finditer(text or ""):
        if match.start() > cursor:
            _set_run_font(p.add_run(text[cursor:match.start()]), size_pt=11, color=COLOR_BODY)
        code = match.group(1)
        number = citation_context["number_map"].get(code)
        shown = f"[{number}]" if code in citation_context["valid_codes"] and number else "[引用异常]"
        run = p.add_run(shown)
        _set_run_font(run, size_pt=8.5, bold=True, color=RGBColor(0x1D, 0x4E, 0x89))
        run.font.superscript = True
        cursor = match.end()
    if cursor < len(text or ""):
        _set_run_font(p.add_run(text[cursor:]), size_pt=11, color=COLOR_BODY)


def _add_report_metadata(doc: Document, report: IndustryReport, grounded: bool) -> None:
    _add_heading(doc, "报告信息", level=2)
    _add_meta_line(doc, "报告版本", f"v{report.version}")
    _add_meta_line(doc, "生成模式", "证据约束（grounded）" if grounded else "传统生成（legacy）")
    _add_meta_line(doc, "生成时间", report.updated_at.strftime("%Y-%m-%d %H:%M") if report.updated_at else "未记录")
    if grounded:
        _add_meta_line(doc, "Prompt版本", report.prompt_version or "未记录")
        _add_meta_line(doc, "引用验证", report.citation_validation_status or "未记录")
        _add_meta_line(doc, "晋升方式", report.promotion_type or "未记录")
        if report.promotion_note:
            _add_meta_line(doc, "审批备注", report.promotion_note)
    else:
        p = doc.add_paragraph()
        run = p.add_run("本报告不是证据约束报告，正文没有经过逐条引用绑定与导出前证据复核。")
        _set_run_font(run, size_pt=10.5, bold=True, color=RGBColor(0x9A, 0x34, 0x12))


def _display_url(value: str) -> str:
    # Zero-width opportunities let Word wrap long URLs without changing the target shown to users.
    return value.replace("/", "/\u200b").replace("?", "?\u200b").replace("&", "&\u200b")


def _add_grounded_appendices(doc: Document, context: dict[str, Any]) -> None:
    doc.add_page_break()
    _add_heading(doc, "附录一：引用与来源", level=1)
    for item in context["citations"]:
        _add_heading(doc, f"[{item['display_number']}] {item['source_name']}", level=2)
        if item.get("source_publisher"):
            _add_meta_line(doc, "发布机构", item["source_publisher"])
        _add_meta_line(doc, "来源类型", item.get("source_origin") or "未记录")
        _add_meta_line(doc, "证据等级", item.get("evidence_grade") or "未记录")
        _add_meta_line(doc, "原文位置", item.get("locator") or "未记录")
        if item.get("published_at"):
            _add_meta_line(doc, "发布日期", item["published_at"])
        if item.get("retrieved_at"):
            _add_meta_line(doc, "获取时间", item["retrieved_at"])
        if item.get("url"):
            _add_meta_line(doc, "网页地址", _display_url(item["url"]))
        _add_meta_line(doc, "原文摘录", item.get("original_quote") or "未记录")

    doc.add_page_break()
    _add_heading(doc, "附录二：冲突与限制", level=1)
    reported_conflicts = [str(item) for item in context.get("unresolved_conflicts") or []]
    for item in reported_conflicts:
        _add_body(doc, f"• {item}")
    seen: set[str] = set()
    for item in context["citations"]:
        for conflict in item.get("related_conflicts") or []:
            code = str(conflict.get("conflict_code") or "")
            if not code or code in seen:
                continue
            seen.add(code)
            _add_heading(doc, f"{code} · {conflict.get('status') or '未记录'}", level=2)
            _add_body(doc, str(conflict.get("description") or ""))
            if conflict.get("resolution_note"):
                _add_meta_line(doc, "处理说明", str(conflict["resolution_note"]))
    if not seen and not reported_conflicts:
        _add_body(doc, "当前引用未关联已登记的跨信源冲突。")
    limitations = list(context.get("limitations") or [])
    source_limitations = sorted({
        limitation
        for item in context["citations"] for limitation in item.get("limitations") or []
    })
    _add_heading(doc, "资料与方法限制", level=2)
    for limitation in limitations + source_limitations:
        _add_body(doc, f"• {limitation}")
    if not limitations and not source_limitations:
        _add_body(doc, "未登记额外限制。")

    doc.add_page_break()
    _add_heading(doc, "附录三：证据覆盖", level=1)
    coverage = context.get("coverage") or {}
    labels = {
        "verified_evidence_count": "可用验证证据数",
        "independent_source_count": "独立来源数",
        "partial_text_count": "部分正文证据数",
        "unresolved_conflict_count": "未解决冲突数",
        "resolved_conflict_count": "已处理冲突数",
        "network_lead_count": "网络线索数",
        "missing_topic_count": "缺失主题数",
        "factual_citation_coverage": "事实句引用覆盖率",
    }
    shown = False
    for key, label in labels.items():
        if key in coverage:
            value = coverage[key]
            if key == "factual_citation_coverage" and isinstance(value, (int, float)):
                value = f"{value:.1%}"
            _add_meta_line(doc, label, str(value))
            shown = True
    missing_topics = coverage.get("missing_topics") or []
    if missing_topics:
        _add_meta_line(doc, "缺失主题", "、".join(str(item) for item in missing_topics))
        shown = True
    if not shown:
        _add_body(doc, "没有可展示的覆盖统计。")


def build_industry_report_docx(
    report: IndustryReport, citation_context: Optional[dict[str, Any]] = None,
) -> Document:
    import json

    doc = Document()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run(f"{report.industry_name} — 行业分析报告")
    _set_run_font(tr, size_pt=20, bold=True, color=RGBColor(0x0F, 0x17, 0x2A))

    if report.company_name:
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sub.add_run(f"分析对象: {report.company_name}")
        _set_run_font(sr, size_pt=11, color=RGBColor(0x55, 0x55, 0x55))

    grounded = report.generation_mode == "grounded"
    if grounded and citation_context is None:
        raise ValueError("grounded_report_requires_validated_citation_context")
    _add_report_metadata(doc, report, grounded)

    payload = {}
    if report.report_json:
        try:
            payload = json.loads(report.report_json)
        except json.JSONDecodeError:
            payload = {}

    if payload.get("summary"):
        _add_heading(doc, "执行摘要", level=1)
        if grounded:
            _add_cited_body(doc, str(payload["summary"]), citation_context)
        else:
            _add_body(doc, str(payload["summary"]))
    for sec in payload.get("sections") or []:
        if isinstance(sec, dict):
            _add_heading(doc, str(sec.get("heading", "")), level=2)
            if grounded:
                _add_cited_body(doc, str(sec.get("content", "")), citation_context)
            else:
                _add_body(doc, str(sec.get("content", "")))
    if payload.get("risk_outlook"):
        _add_heading(doc, "风险展望", level=1)
        if grounded:
            _add_cited_body(doc, str(payload["risk_outlook"]), citation_context)
        else:
            _add_body(doc, str(payload["risk_outlook"]))

    if payload.get("key_metrics"):
        _add_heading(doc, "关键指标", level=1)
        for metric in payload["key_metrics"]:
            if not isinstance(metric, dict):
                continue
            text = f"{metric.get('name', '')}：{metric.get('value', '')}"
            code = str(metric.get("evidence_code") or "")
            if grounded and code:
                text += f"[{code}]"
                _add_cited_body(doc, text, citation_context)
            else:
                _add_body(doc, text)

    if report.chart_specs:
        try:
            charts = json.loads(report.chart_specs)
        except json.JSONDecodeError:
            charts = []
        for idx, item in enumerate(charts):
            opt = item.get("option") if isinstance(item, dict) else None
            if not opt:
                continue
            spec = {
                "type": (opt.get("series") or [{}])[0].get("type", "bar"),
                "title": (opt.get("title") or {}).get("text", "图表"),
                "labels": (opt.get("xAxis") or {}).get("data", []),
                "series": [{"name": s.get("name", ""), "data": s.get("data", [])} for s in opt.get("series", [])],
            }
            img_path = render_chart_png(spec, Path(f"data/exports/charts/industry_{report.id}_{idx}.png"))
            if img_path and img_path.is_file():
                doc.add_picture(str(img_path), width=Inches(5.5))

    if grounded:
        _add_grounded_appendices(doc, citation_context)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("— 内部资料，未经许可不得对外传播 —")
    _set_run_font(fr, size_pt=9, color=RGBColor(0x77, 0x77, 0x77))
    return doc


def export_industry_report_to_path(
    report: IndustryReport, output_path: Path,
    citation_context: Optional[dict[str, Any]] = None,
) -> Path:
    doc = build_industry_report_docx(report, citation_context=citation_context)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def export_daily_report_to_path(
    entries: list[DailyRiskEntry],
    report_date: date,
    output_path: Path,
    institution_name: str = "风险管理部",
    module_codes: Optional[list[str]] = None,
) -> Path:
    doc = build_daily_report_docx(
        entries, report_date, institution_name, module_codes=module_codes
    )
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def export_daily_report_to_bytes(
    entries: list[DailyRiskEntry],
    report_date: date,
    institution_name: str = "风险管理部",
    module_codes: Optional[list[str]] = None,
) -> bytes:
    doc = build_daily_report_docx(
        entries, report_date, institution_name, module_codes=module_codes
    )
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_entity_assessment_docx(
    entity: TargetEntity,
    *,
    report_date: date,
    risks: Iterable[EntityRisk],
    credit_logs: Iterable[CreditUpdate],
    institution_name: str = "风险管理部",
) -> Document:
    """生成《企业主体风险评估简报》。"""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)

    display = entity.display_name or entity.name

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run("企业主体风险评估简报")
    _set_run_font(tr, size_pt=22, bold=True, color=RGBColor(0x0F, 0x17, 0x2A))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(f"{institution_name} | 评估日期 {report_date.isoformat()}")
    _set_run_font(sr, size_pt=11, color=RGBColor(0x55, 0x55, 0x55))

    doc.add_paragraph()
    _add_heading(doc, "一、主体概况", level=1)
    _add_meta_line(doc, "主体名称", display)
    if entity.industry:
        _add_meta_line(doc, "所属行业", entity.industry)
    if entity.region:
        _add_meta_line(doc, "区域", entity.region)
    _add_meta_line(doc, "当前授信等级", entity.credit_level or "正常")
    _add_meta_line(doc, "监控状态", entity.monitor_status or "active")

    _add_heading(doc, "二、授信变更历史", level=1)
    logs = list(credit_logs)
    if not logs:
        _add_body(doc, "暂无授信变更记录。", indent=True)
    else:
        for idx, log in enumerate(logs, start=1):
            when = log.created_at.isoformat(sep=" ", timespec="minutes") if log.created_at else ""
            _add_heading(doc, f"{idx}. {log.previous_level} → {log.new_level}", level=2)
            if when:
                _add_meta_line(doc, "变更时间", when)
            _add_meta_line(doc, "调级原因", log.reason or "未注明")

    _add_heading(doc, "三、近期风险事件与 AI 摘要", level=1)
    risk_list = list(risks)
    if not risk_list:
        _add_body(doc, "近期暂无风险事件。", indent=True)
    else:
        for idx, e in enumerate(risk_list, start=1):
            _add_heading(doc, f"{idx}. {e.title}", level=2)
            _add_meta_line(doc, "发生/报告日期", e.report_date.isoformat())
            _add_meta_line(doc, "风险等级", e.risk_level)
            if e.risk_category:
                _add_meta_line(doc, "风险类别", e.risk_category)
            _add_meta_line(doc, "核心摘要", e.summary or "")
            if e.impact_analysis:
                _add_meta_line(doc, "AI 风险影响", e.impact_analysis)
            if e.source_url:
                _add_meta_line(doc, "数据来源", e.source_url)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("— 内部资料，未经许可不得对外传播 —")
    _set_run_font(fr, size_pt=9, color=RGBColor(0x77, 0x77, 0x77))
    return doc


def export_entity_assessment_to_path(
    entity: TargetEntity,
    *,
    report_date: date,
    risks: list[EntityRisk],
    credit_logs: list[CreditUpdate],
    output_path: Path,
    institution_name: str = "风险管理部",
) -> Path:
    doc = build_entity_assessment_docx(
        entity,
        report_date=report_date,
        risks=risks,
        credit_logs=credit_logs,
        institution_name=institution_name,
    )
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
